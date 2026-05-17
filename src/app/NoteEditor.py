"""Note editor page for creating and editing campaign notes."""

import io
import json
import os

import pandas as pd
import streamlit as st
from docx import Document as _DocxDocument
from html.parser import HTMLParser
from streamlit_lottie import st_lottie

from ..utils import DatabaseHandler
from ..utils.TextEditorHandler import TextEditorHandler
from ..utils import paths
from .. import app_config


# ---------------------------------------------------------------------------
# Module-level utilities — pure functions with no Streamlit dependency
# ---------------------------------------------------------------------------

class _HTMLTextExtractor(HTMLParser):
    _BLOCK_TAGS = frozenset({"p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr"})

    def __init__(self):
        super().__init__()
        self._parts = []

    def handle_starttag(self, tag, attrs):
        if tag in self._BLOCK_TAGS:
            self._parts.append("\n")

    def handle_data(self, data):
        self._parts.append(data)

    def get_text(self):
        return "".join(self._parts).strip()


def strip_html(html_content: str) -> str:
    """Return plain text with HTML tags removed."""
    extractor = _HTMLTextExtractor()
    extractor.feed(html_content or "")
    return extractor.get_text()


def load_editor_notes(filepath: str) -> str:
    """Return saved notes from disk, or empty string if unavailable."""
    if os.path.isfile(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            pass
    return ""


def save_editor_notes(filepath: str, content: str) -> None:
    """Persist notes to disk, creating parent directories as needed."""
    dir_path = os.path.dirname(filepath)
    if dir_path:
        os.makedirs(dir_path, exist_ok=True)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content or "")


def raw_notes_to_text(raw_notes_path: str) -> str:
    """Convert data/raw_notes.json (DataFrame JSON) to plain text."""
    if not os.path.isfile(raw_notes_path):
        return ""
    try:
        df = pd.read_json(raw_notes_path)
    except Exception:
        return ""
    parts = []
    for _, row in df.iterrows():
        date = str(row.get("Date", "") or "")
        title = str(row.get("Title", "") or "")
        contents = str(row.get("Contents", "") or "")
        header = date if date.lower() not in ("", "unknown date", "nan") else title
        if header:
            parts.append(header)
        if contents:
            parts.append(contents)
        parts.append("")
    return "\n".join(parts).strip()


_DEFAULT_EDITOR_CONFIG = {"font_family": "Georgia", "font_size": 16}

_FONT_OPTIONS = ["Georgia", "Arial", "Courier New", "Times New Roman", "Palatino Linotype"]


def load_editor_config(filepath: str) -> dict:
    """Return editor config from disk, or defaults if unavailable."""
    if os.path.isfile(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return dict(_DEFAULT_EDITOR_CONFIG)


def save_editor_config(filepath: str, config: dict) -> None:
    """Persist editor config to disk, creating parent directories as needed."""
    dir_path = os.path.dirname(filepath)
    if dir_path:
        os.makedirs(dir_path, exist_ok=True)
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(config, f)


def build_txt_content(content: str) -> str:
    """Return plain text content for TXT export."""
    return content or ""


def build_docx_bytes(content: str) -> bytes:
    """Convert plain text to DOCX bytes for download."""
    doc = _DocxDocument()
    for line in (content or "").splitlines():
        stripped = line.strip()
        if stripped:
            doc.add_paragraph(stripped)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _EditorDocument:
    """Wraps plain text as a file-like object compatible with DatabaseHandler."""

    def __init__(self, text_content: str):
        self.name = "editor_notes.txt"
        self._data = text_content.encode("utf-8")

    def getvalue(self) -> bytes:
        return self._data

    def read(self) -> bytes:
        return self._data


# ---------------------------------------------------------------------------
# NoteEditor page class
# ---------------------------------------------------------------------------

class NoteEditor:
    _NOTES_FILE = "data/editor_notes.txt"
    _RAW_NOTES_FILE = "data/raw_notes.json"
    _CAMPAIGN_SUMMARY_FILE = "data/campaign_summary.json"
    _CONFIG_FILE = "data/editor_config.json"

    def __init__(self):
        if "databasehandler" not in st.session_state:
            st.session_state.databasehandler = DatabaseHandler.DatabaseHandler()
        self.databasehandler = st.session_state.databasehandler
        self._editor = TextEditorHandler()

        if app_config.is_remote():
            uid = st.session_state.get("auth_user_id")
            paths.ensure_data_root(uid)
            self._NOTES_FILE = paths.editor_notes_file(uid)
            self._RAW_NOTES_FILE = paths.raw_notes_file(uid)
            self._CAMPAIGN_SUMMARY_FILE = paths.summary_file(uid)
            self._CONFIG_FILE = paths.editor_config_file(uid)
            self._DATABASEDIR = paths.database_dir(uid)
        else:
            self._DATABASEDIR = DatabaseHandler.DATABASE_DIR

    def __init_state_variables(self):
        if "editor_content" not in st.session_state:
            saved = load_editor_notes(self._NOTES_FILE)
            if not saved and os.path.isfile(self._RAW_NOTES_FILE):
                saved = raw_notes_to_text(self._RAW_NOTES_FILE)
            st.session_state.editor_content = saved
        if "editor_key" not in st.session_state:
            st.session_state.editor_key = 0
        if "editor_font_family" not in st.session_state:
            config = load_editor_config(self._CONFIG_FILE)
            st.session_state.editor_font_family = config.get("font_family", _DEFAULT_EDITOR_CONFIG["font_family"])
            st.session_state.editor_font_size = config.get("font_size", _DEFAULT_EDITOR_CONFIG["font_size"])

    def __import_uploaded_notes(self):
        text = raw_notes_to_text(self._RAW_NOTES_FILE)
        if text:
            st.session_state.editor_content = text
            save_editor_notes(self._NOTES_FILE, text)
            st.session_state.editor_key += 1

    def __vectorize_notes(self):
        content = st.session_state.editor_content
        if not content or not content.strip():
            st.error("No content in the editor to vectorize.")
            return

        self.databasehandler.clear_database(self._DATABASEDIR)
        for stale in (self._RAW_NOTES_FILE, self._CAMPAIGN_SUMMARY_FILE):
            if os.path.isfile(stale):
                os.remove(stale)

        self.databasehandler.create_retrival_artifacts(self._DATABASEDIR)

        doc = _EditorDocument(content)

        try:
            with open("assets/Magical_Effect_Loading.json", "r", errors="ignore") as f:
                magic_loader = json.load(f)
        except Exception:
            magic_loader = None

        animation_slot = st.empty()
        progress_slot = st.empty()

        if magic_loader:
            with animation_slot.container():
                st_lottie(magic_loader, height=200, key="editor_vectorize_spinner")

        progress_bar = progress_slot.progress(0, text="Vectorizing notes...")
        gen = self.databasehandler.generate_database(doc, self._DATABASEDIR)
        return_code = None

        while True:
            try:
                progress = next(gen)
                progress_bar.progress(progress / 100, text=f"Vectorizing notes... {progress:.1f}%")
            except StopIteration as e:
                return_code = e.value
                break

        animation_slot.empty()
        progress_slot.empty()

        if return_code:
            if self.databasehandler.last_processed_df is not None:
                os.makedirs(os.path.dirname(self._RAW_NOTES_FILE) or ".", exist_ok=True)
                self.databasehandler.last_processed_df.to_json(self._RAW_NOTES_FILE)
            st.toast("📜 Notes vectorized successfully!", icon="🧙‍♂️")
        else:
            st.error("Vectorization failed. Ensure the editor contains valid content.")

    def __render_sidebar(self):
        with st.sidebar:
            st.header("📜 Note Options")

            raw_notes_exist = os.path.isfile(self._RAW_NOTES_FILE)
            if st.button(
                "📥 Load from Uploaded Notes",
                disabled=not raw_notes_exist,
                help="No uploaded notes found." if not raw_notes_exist else "Replace editor content with notes from the last file upload.",
                use_container_width=True,
            ):
                self.__import_uploaded_notes()
                st.rerun()

            st.divider()

            st.subheader("🔤 Text Formatting")
            current_font = st.session_state.get("editor_font_family", _DEFAULT_EDITOR_CONFIG["font_family"])
            current_size = st.session_state.get("editor_font_size", _DEFAULT_EDITOR_CONFIG["font_size"])
            font_idx = _FONT_OPTIONS.index(current_font) if current_font in _FONT_OPTIONS else 0
            new_font = st.selectbox("Font", _FONT_OPTIONS, index=font_idx)
            new_size = st.slider("Size (px)", min_value=10, max_value=28, value=current_size, step=2)
            if new_font != current_font or new_size != current_size:
                st.session_state.editor_font_family = new_font
                st.session_state.editor_font_size = new_size
                save_editor_config(self._CONFIG_FILE, {"font_family": new_font, "font_size": new_size})

            st.divider()

            st.subheader("🧠 Vectorize")
            has_content = bool(st.session_state.get("editor_content", "").strip())
            if st.button(
                "⚡ Vectorize Notes",
                type="primary",
                use_container_width=True,
                disabled=not has_content,
                help="No content to vectorize." if not has_content else "Clear the database and re-vectorize the current editor notes.",
            ):
                st.session_state._do_vectorize = True
                st.rerun()

            st.divider()

            st.subheader("📤 Export")
            content = st.session_state.get("editor_content", "")
            st.download_button(
                label="📄 Export as TXT",
                data=build_txt_content(content),
                file_name="campaign_notes.txt",
                mime="text/plain",
                use_container_width=True,
            )
            st.download_button(
                label="📝 Export as DOCX",
                data=build_docx_bytes(content),
                file_name="campaign_notes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

    def __render_editor(self):
        new_content = self._editor.render(
            key=f"note_editor_{st.session_state.editor_key}",
            initial_value=st.session_state.get("editor_content", ""),
            height=600,
            font_family=st.session_state.get("editor_font_family", _DEFAULT_EDITOR_CONFIG["font_family"]),
            font_size=st.session_state.get("editor_font_size", _DEFAULT_EDITOR_CONFIG["font_size"]),
        )
        if new_content != st.session_state.editor_content:
            st.session_state.editor_content = new_content
            save_editor_notes(self._NOTES_FILE, new_content)

    def run(self):
        st.title("📝 Note Editor")
        st.info(
            "Create and edit your campaign notes here. "
            "Notes are saved automatically and persist between sessions."
        )

        self.__init_state_variables()
        self.__render_sidebar()

        if st.session_state.pop("_do_vectorize", False):
            self.__vectorize_notes()

        self.__render_editor()
